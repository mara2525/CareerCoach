"""
Cover Letter Generator for Mara Jorgensen

Usage:
    Customize the CONTENT section below for each job application,
    then run: python create_cover_letter.py

Output:
    Creates .docx and .pdf files in the output/ folder
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants
from datetime import datetime
import os

# =============================================================================
# CONTENT - Customize this section for each job application
# =============================================================================

COMPANY = "Datamundi"
ROLE = "Insurance SME AI Evaluation"
DATE = datetime.now().strftime("%B %d, %Y")  # e.g., "January 26, 2026"

# Attention line (usually "[Role] Hiring Manager" or specific person's title)
ATTN = "AI Evaluation Hiring Team"

# Paragraphs - customize these for each application (3-5 paragraphs max)
PARAGRAPHS = [
    # Opening paragraph
    f"I am excited to apply for the {ROLE} position at {COMPANY}. I bring 10+ years of insurance domain expertise (trucking, health, crop) combined with hands-on AI evaluation and prompt engineering experience.",

    # Body paragraph - Insurance + AI experience
    "At Roger, I led a trucking insurance compliance product from 0-1, building systems that validated carrier insurance against FMCSA requirements and increased ARR by 20%. I have also developed AI applications for health insurance enrollment and crop insurance risk assessment. This gives me direct experience applying LLMs to underwriting and claims workflows.",

    # Body paragraph - AI evaluation expertise
    "I work daily with prompt engineering, evaluation rubrics, and golden response development. I validate AI outputs for accuracy, bias, and compliance, particularly around sensitive data like PHI. I collaborate effectively with data science teams while providing the insurance context needed to improve model performance.",

    # Closing paragraph
    "I would welcome the opportunity to discuss how my combined insurance and AI evaluation experience can contribute to Datamundi's client success. Thank you for your consideration.",
]

# =============================================================================
# STYLING - Generally don't need to modify
# =============================================================================

DARK_GRAY = RGBColor(0x3D, 0x3D, 0x3D)

def add_hyperlink(paragraph, text, url, color=DARK_GRAY, font_size=Pt(11), font_name='Gill Sans MT'):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '%02X%02X%02X' % (color[0], color[1], color[2]))
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(szCs)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def create_cover_letter():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Gill Sans MT'
    style.font.size = Pt(11)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ==========  HEADER ==========
    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Mara Jorgensen")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Gill Sans MT'
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

    # Email
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_hyperlink(p, "mara.jorgensen@gmail.com", "mailto:mara.jorgensen@gmail.com")
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

    # Phone
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("712-898-2341")
    run.font.size = Pt(11)
    run.font.name = 'Gill Sans MT'
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

    # LinkedIn
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_hyperlink(p, "LinkedIn", "https://www.linkedin.com/in/mara-jorgensen/")
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(DATE)
    run.font.size = Pt(11)
    run.font.name = 'Gill Sans MT'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.0

    # ==========  ATTENTION LINE ==========
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"ATTN: {ATTN}")
    run.font.size = Pt(11)
    run.font.name = 'Gill Sans MT'
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.0

    # ==========  GREETING ==========
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Hi,")
    run.font.size = Pt(11)
    run.font.name = 'Gill Sans MT'
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.0

    # ==========  BODY PARAGRAPHS ==========
    for para_text in PARAGRAPHS:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(para_text)
        run.font.size = Pt(11)
        run.font.name = 'Gill Sans MT'
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.15

    # ==========  SIGN-OFF ==========
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Sincerely,")
    run.font.size = Pt(11)
    run.font.name = 'Gill Sans MT'
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Mara Jorgensen")
    run.font.size = Pt(11)
    run.font.name = 'Gill Sans MT'
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

    return doc

def main():
    # Create cover letter
    doc = create_cover_letter()

    # Generate filename
    company_clean = COMPANY.replace(" ", "_").replace("/", "-")
    date_str = datetime.now().strftime("%Y-%m")
    base_filename = f"Jorgensen Cover Letter - {COMPANY} {ROLE} {date_str}"

    # Output paths
    output_dir = os.path.dirname(os.path.abspath(__file__))
    if not output_dir.endswith("output"):
        output_dir = os.path.join(output_dir, "output")

    docx_path = os.path.join(output_dir, f"{base_filename}.docx")
    pdf_path = os.path.join(output_dir, f"{base_filename}.pdf")

    # Save docx
    doc.save(docx_path)
    print(f"Created: {docx_path}")

    # Convert to PDF using Microsoft Word COM
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc_obj = word.Documents.Open(docx_path)
        doc_obj.SaveAs(pdf_path, FileFormat=17)  # 17 = wdFormatPDF
        doc_obj.Close()
        word.Quit()
        print(f"Created: {pdf_path}")
    except Exception as e:
        print(f"PDF conversion failed: {e}")
        print("The .docx file was created successfully. Export to PDF manually from Word.")

if __name__ == "__main__":
    main()
