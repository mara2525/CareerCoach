from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants

# Colors
DUSTY_BLUE = RGBColor(0x53, 0x76, 0x85)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x3D, 0x3D, 0x3D)

def add_hyperlink(paragraph, text, url, color=DARK_GRAY, font_size=Pt(10), font_name='Calibri'):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # Color
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '%02X%02X%02X' % (color[0], color[1], color[2]))
    rPr.append(c)
    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    # Font size
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(szCs)
    # Font name
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)
    new_run.append(rPr)
    # Text
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Set narrow margins
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

def add_section_header(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = DUSTY_BLUE
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '537685')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0

def add_role_bullet(lead_in, description):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(lead_in + ": ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run = p.add_run(description)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0

def add_simple_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0

# ========== HEADER ==========
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MARA JORGENSEN")
run.bold = True
run.font.size = Pt(22)
run.font.name = 'Gill Sans MT'
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AI Product Leader | Enterprise Platforms & Agentic AI")
run.italic = True
run.font.size = Pt(11)
run.font.color.rgb = DUSTY_BLUE
run.font.name = 'Gill Sans MT'
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("+1-712-898-2341 | ")
run.font.size = Pt(10)
run.font.color.rgb = DARK_GRAY
run.font.name = 'Calibri'
add_hyperlink(p, "mara@strategic-corp.com", "mailto:mara@strategic-corp.com")
run = p.add_run(" | ")
run.font.size = Pt(10)
run.font.color.rgb = DARK_GRAY
run.font.name = 'Calibri'
add_hyperlink(p, "LinkedIn", "https://www.linkedin.com/in/mara-jorgensen/")
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.line_spacing = 1.0

# ========== SUMMARY ==========
add_section_header("Summary")
p = doc.add_paragraph()
run = p.add_run("Technical Product Leader with 20+ years building enterprise platforms, now focused on AI-driven products. Expert at taking products from 0-1, leading cross-functional teams through ambiguity to deliver customer value. I combine hands-on technical depth (Python, APIs, LLM orchestration) with strategic product vision to build AI agents and platforms that solve real business problems. Passionate about customer obsession: deeply understanding user needs and iterating rapidly to deliver excellent experiences.")
run.font.size = Pt(10.5)
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.line_spacing = 1.0

# ========== SKILLS ==========
add_section_header("Skills")

p = doc.add_paragraph()
run = p.add_run("AI & Agentic Product Leadership")
run.bold = True
run.font.size = Pt(10.5)
run.font.name = 'Gill Sans MT'
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing = 1.0

for skill in ["Agentic Workflows & LLM Orchestration", "AI Roadmapping (90-day and beyond)", "Rapid Pilot-to-Production Cycles", "AI Product Lifecycle Ownership", "Responsible AI Posture"]:
    add_simple_bullet(skill)

p = doc.add_paragraph()
run = p.add_run("Product Strategy & Execution")
run.bold = True
run.font.size = Pt(10.5)
run.font.name = 'Gill Sans MT'
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing = 1.0

for skill in ["0-1 Product Launches", "Enterprise Product Strategy", "P&L Ownership", "User Discovery, A/B Testing, Customer Flows", "MVP Definition & Prioritization"]:
    add_simple_bullet(skill)

p = doc.add_paragraph()
run = p.add_run("Technical & Platform")
run.bold = True
run.font.size = Pt(10.5)
run.font.name = 'Gill Sans MT'
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing = 1.0

for skill in ["APIs, SDKs, & Developer Tooling", "Python, SQL, Integration Architecture", "AI Tools: OpenAI, Claude, Cursor, v0, agents, workflows", "Cross-functional Engineering Leadership"]:
    add_simple_bullet(skill)

# ========== PROFESSIONAL EXPERIENCE ==========
add_section_header("Professional Experience")

# Strategic Consulting Corp
p = doc.add_paragraph()
run = p.add_run("CEO, AI & Technical Implementation Strategist")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Gill Sans MT'
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0

p = doc.add_paragraph()
run = p.add_run("Strategic Consulting Corp.")
run.font.color.rgb = DUSTY_BLUE
run.font.size = Pt(10)
run.font.name = 'Gill Sans MT'
run = p.add_run(" | 1/2020 – Present")
run.font.size = Pt(10)
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing = 1.0

add_role_bullet("AI Framework Development", "Developed and executes a 3-phase framework to move enterprise clients from concept to functional AI pilot, validating core business assumptions quickly")
add_role_bullet("Agentic Engineering", "Leverages a \"next-gen\" tech stack including AI coding assistants and agentic workflows to deliver functional pilots at a fraction of traditional engineering time and cost")
add_role_bullet("Executive Advisory", "Consults with executives to identify AI use-cases that provide immediate value (ROI, efficiency gains, customer experience improvements)")
add_role_bullet("Technical Bridge", "Represents engineering and data needs in executive settings, translating complex technical concepts for business stakeholders")

# Kontango
p = doc.add_paragraph()
run = p.add_run("Head of Freight Product & Customer Operations")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Gill Sans MT'
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0

p = doc.add_paragraph()
run = p.add_run("Kontango")
run.font.color.rgb = DUSTY_BLUE
run.font.size = Pt(10)
run.font.name = 'Gill Sans MT'
run = p.add_run(" | AI-powered chemicals commodity trading platform | 4/2023 – 12/2024")
run.font.size = Pt(10)
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing = 1.0

add_role_bullet("ML-Driven Product", "Owned end-to-end product roadmap for the freight platform, leading the integration of predictive ML models with the core trading application")
add_role_bullet("Customer Obsession", "Drove rigorous discovery research, A/B testing, and customer journey analysis to define MVP concepts that maximized ROI for freight users")
add_role_bullet("Cross-Functional Leadership", "Led engineering, data science, and commercial teams to rapidly ship features; served as primary customer-facing advisor to executive stakeholders")
add_role_bullet("Rapid Prototyping", "Leveraged AI coding assistants to rapidly engineer and deploy solutions, accelerating the feedback loop with customers")

# Roger
p = doc.add_paragraph()
run = p.add_run("General Manager, Product Management [Integrations & Compliance]")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Gill Sans MT'
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0

p = doc.add_paragraph()
run = p.add_run("Roger (JV of Cargill, Andersons, Scoular, Koch)")
run.font.color.rgb = DUSTY_BLUE
run.font.size = Pt(10)
run.font.name = 'Gill Sans MT'
run = p.add_run(" | Enterprise logistics platform | 3/2020 – 3/2023")
run.font.size = Pt(10)
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing = 1.0

add_role_bullet("0-1 Platform Launch", "Owned P&L for two product lines (Integrations and Compliance); led the design and launch of onRamp, a secure B2B onboarding platform—increased revenue by ~20% within 12 months")
add_role_bullet("Developer Platform", "Spearheaded iPaaS selection and implementation to scale B2B integrations, conducting market analysis and securing executive buy-in")
add_role_bullet("Continuous Discovery", "Led discovery sessions with enterprise stakeholders to translate user needs into detailed technical requirements and user stories")
add_role_bullet("Cross-Functional Execution", "Partnered with engineering, design, and legal to ship compliant features; led onshore and offshore engineering teams")

# Fast Solutions
p = doc.add_paragraph()
run = p.add_run("Partner, Director of Business Solutions")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Gill Sans MT'
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0

p = doc.add_paragraph()
run = p.add_run("Fast Solutions")
run.font.color.rgb = DUSTY_BLUE
run.font.size = Pt(10)
run.font.name = 'Gill Sans MT'
run = p.add_run(" | Enterprise software consultancy | 6/2013 – 1/2020")
run.font.size = Pt(10)
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing = 1.0

add_role_bullet("Technical Leadership", "Led product and engineering teams in designing and delivering complex integrations and platform implementations for enterprise clients")
add_role_bullet("User-Centered Design", "Led user research and stakeholder interviews to inform product roadmaps; managed full product lifecycle from requirements to release")
add_role_bullet("Enterprise Transformation", "Drove digital transformation by creating multi-year product roadmaps for integration and modernization projects")

# Tyson Foods
p = doc.add_paragraph()
run = p.add_run("Business Analyst/Software Engineer")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Gill Sans MT'
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0

p = doc.add_paragraph()
run = p.add_run("Tyson Foods")
run.font.color.rgb = DUSTY_BLUE
run.font.size = Pt(10)
run.font.name = 'Gill Sans MT'
run = p.add_run(" | 7/2010 – 6/2013")
run.font.size = Pt(10)
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing = 1.0

add_simple_bullet("Acted as key liaison between business stakeholders across sales, manufacturing, and accounting")
add_simple_bullet("Developed integrations that enhanced data flow between critical business systems")

# CF Industries
p = doc.add_paragraph()
run = p.add_run("Software Developer")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Gill Sans MT'
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0

p = doc.add_paragraph()
run = p.add_run("CF Industries")
run.font.color.rgb = DUSTY_BLUE
run.font.size = Pt(10)
run.font.name = 'Gill Sans MT'
run = p.add_run(" | 4/2001 – 6/2010")
run.font.size = Pt(10)
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing = 1.0

add_simple_bullet("Engineered robust software solutions and integrations for enterprise manufacturing applications")

# ========== EDUCATION ==========
add_section_header("Education")

p = doc.add_paragraph()
run = p.add_run("Masters in Management")
run.bold = True
run.font.size = Pt(10.5)
run.font.name = 'Calibri'
run = p.add_run(" | Bellevue University, Bellevue, Nebraska")
run.font.size = Pt(10.5)
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing = 1.0

p = doc.add_paragraph()
run = p.add_run("BS in Computer Science and Mathematics")
run.bold = True
run.font.size = Pt(10.5)
run.font.name = 'Calibri'
run = p.add_run(" | Morningside University, Sioux City, Iowa")
run.font.size = Pt(10.5)
run.font.name = 'Calibri'

# Save
output_path = r"g:\My Drive\GitHub\ExecutiveAssistant\Resume\output\Jorgensen Resume - Sierra PM 2026-01.docx"
doc.save(output_path)
print(f"Created: {output_path}")
