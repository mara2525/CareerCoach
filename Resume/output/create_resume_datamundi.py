from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants

# Colors
DUSTY_BLUE = RGBColor(0x53, 0x76, 0x85)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x3D, 0x3D, 0x3D)

def add_hyperlink(paragraph, text, url, color=DARK_GRAY, font_size=Pt(10), font_name='Calibri'):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # Color
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '%02X%02X%02X' % (color[0], color[1], color[2]))
    rPr.append(c)
    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    # Font size
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(szCs)
    # Font name
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)
    new_run.append(rPr)
    # Text
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Set narrow margins
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

def add_section_header(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = DUSTY_BLUE
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '537685')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0

def add_role_bullet(lead_in, description):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(lead_in + ": ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run = p.add_run(description)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0

def add_simple_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0

# ========== HEADER ==========
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MARA JORGENSEN")
run.bold = True
run.font.size = Pt(22)
run.font.name = 'Gill Sans MT'
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Insurance Domain Expert | AI Evaluation & Prompt Engineering")
run.italic = True
run.font.size = Pt(11)
run.font.color.rgb = DUSTY_BLUE
run.font.name = 'Gill Sans MT'
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("+1-712-898-2341 | ")
run.font.size = Pt(10)
run.font.color.rgb = DARK_GRAY
run.font.name = 'Calibri'
add_hyperlink(p, "mara.jorgensen@gmail.com", "mailto:mara.jorgensen@gmail.com")
run = p.add_run(" | ")
run.font.size = Pt(10)
run.font.color.rgb = DARK_GRAY
run.font.name = 'Calibri'
add_hyperlink(p, "LinkedIn", "https://www.linkedin.com/in/mara-jorgensen/")
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.line_spacing = 1.0

# ========== SUMMARY ==========
add_section_header("Summary")
p = doc.add_paragraph()
run = p.add_run("Insurance domain expert and AI expert with 10+ years of experience spanning trucking/commercial auto, health insurance enrollment, and crop insurance. Hands-on AI/LLM expertise in prompt engineering, model evaluation, and rubric development. Led a trucking insurance compliance product from 0-1 that increased ARR by 20%. I bridge deep insurance knowledge (policy lifecycle, risk assessment, regulatory compliance) with technical AI capabilities to evaluate and improve AI model outputs for accuracy, bias, and compliance. Experienced collaborator with cross-functional teams including data science, engineering, and product.")
run.font.size = Pt(10.5)
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.line_spacing = 1.0

# ========== SKILLS ==========
add_section_header("Skills")

p = doc.add_paragraph()
run = p.add_run("Insurance Domain Expertise")
run.bold = True
run.font.size = Pt(10.5)
run.font.name = 'Gill Sans MT'
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing = 1.0

for skill in ["Trucking/Commercial Auto Insurance & FMCSA Compliance", "Health Insurance Enrollment & PHI Compliance", "Crop Insurance Risk Assessment", "Policy Lifecycle, Retention & Renewal Analysis", "Regulatory Frameworks & Compliance Documentation"]:
    add_simple_bullet(skill)

p = doc.add_paragraph()
run = p.add_run("AI Evaluation & Prompt Engineering")
run.bold = True
run.font.size = Pt(10.5)
run.font.name = 'Gill Sans MT'
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing = 1.0

for skill in ["Prompt Engineering & System Prompt Design", "Evaluation Rubrics & Golden Response Development", "Model Output Validation (Accuracy, Bias, Compliance)", "Data Annotation Guidelines & Quality Assurance", "Explainable AI Framework Design"]:
    add_simple_bullet(skill)

p = doc.add_paragraph()
run = p.add_run("Technical & Collaboration")
run.bold = True
run.font.size = Pt(10.5)
run.font.name = 'Gill Sans MT'
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing = 1.0

for skill in ["AI Tools: Azure AI Foundry, n8n, OpenAI, Claude, LLM Orchestration", "Python, SQL, API Integrations", "Cross-Functional Team Collaboration (Data Science, Engineering, Product)", "Data Analytics, Dashboards & Forecasting"]:
    add_simple_bullet(skill)

# ========== PROFESSIONAL EXPERIENCE ==========
add_section_header("Professional Experience")

# Strategic Consulting Corp
p = doc.add_paragraph()
run = p.add_run("CEO, AI & Technical Implementation Strategist")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Gill Sans MT'
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0

p = doc.add_paragraph()
run = p.add_run("Strategic Consulting Corp.")
run.font.color.rgb = DUSTY_BLUE
run.font.size = Pt(10)
run.font.name = 'Gill Sans MT'
run = p.add_run(" | 1/2020 - Present")
run.font.size = Pt(10)
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing = 1.0

add_role_bullet("AI Application Development", "Build AI/LLM applications for insurance and enterprise clients, including health insurance enrollment and crop insurance risk assessment tools")
add_role_bullet("Prompt Engineering & Evaluation", "Design system prompts, develop evaluation rubrics and golden responses, and validate model outputs against real-world insurance standards and compliance requirements")
add_role_bullet("Model Evaluation", "Test AI outputs for accuracy, bias, and regulatory compliance, with particular focus on sensitive data handling (PHI, PII) in insurance contexts")
add_role_bullet("Cross-Functional Collaboration", "Partner with data science teams to provide insurance domain context for model training and evaluation, ensuring outputs align with industry practices")

# Kontango
p = doc.add_paragraph()
run = p.add_run("Head of Freight Product & Customer Operations")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Gill Sans MT'
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0

p = doc.add_paragraph()
run = p.add_run("Kontango")
run.font.color.rgb = DUSTY_BLUE
run.font.size = Pt(10)
run.font.name = 'Gill Sans MT'
run = p.add_run(" | AI-powered chemicals commodity trading platform | 4/2023 - 12/2024")
run.font.size = Pt(10)
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing = 1.0

add_role_bullet("ML-Driven Product", "Owned end-to-end product roadmap for the freight platform, leading the integration of predictive ML models with the core trading application")
add_role_bullet("Data Analytics & Forecasting", "Built dashboards and forecasting tools for executive stakeholders; used data analytics to drive decision-making and measure model performance")
add_role_bullet("Cross-Functional Leadership", "Led engineering, data science, and commercial teams to rapidly ship features; defined success metrics aligned with business objectives")

# Roger
p = doc.add_paragraph()
run = p.add_run("General Manager, Product Management [Integrations & Compliance]")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Gill Sans MT'
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0

p = doc.add_paragraph()
run = p.add_run("Roger (JV of Cargill, Andersons, Scoular, Koch)")
run.font.color.rgb = DUSTY_BLUE
run.font.size = Pt(10)
run.font.name = 'Gill Sans MT'
run = p.add_run(" | Enterprise logistics platform | 3/2020 - 3/2023")
run.font.size = Pt(10)
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing = 1.0

add_role_bullet("Trucking Insurance Product (0-1)", "Led the compliance product team, building systems to validate carrier insurance compliance in real time; owned the full policy lifecycle from carrier onboarding through certificate verification and renewal tracking")
add_role_bullet("FMCSA Integration", "Integrated with the Federal Motor Carrier Safety Administration to automate insurance documentation verification; required deep understanding of commercial auto policy structures and regulatory frameworks")
add_role_bullet("Revenue Impact", "Increased annual recurring revenue by 20% within 12 months of product launch through insurance compliance features")
add_role_bullet("Enterprise Compliance", "Ensured strict adherence to PII, CCPA, and SOC2 requirements; partnered with legal to ship compliant features")

# Fast Solutions
p = doc.add_paragraph()
run = p.add_run("Partner, Director of Business Solutions")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Gill Sans MT'
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0

p = doc.add_paragraph()
run = p.add_run("Fast Solutions")
run.font.color.rgb = DUSTY_BLUE
run.font.size = Pt(10)
run.font.name = 'Gill Sans MT'
run = p.add_run(" | Enterprise software consultancy | 6/2013 - 1/2020")
run.font.size = Pt(10)
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing = 1.0

add_role_bullet("Technical Leadership", "Led product and engineering teams in designing and delivering complex integrations and platform implementations for enterprise clients")
add_role_bullet("User-Centered Design", "Led user research and stakeholder interviews to inform product roadmaps; managed full product lifecycle from requirements to release")
add_role_bullet("Enterprise Transformation", "Drove digital transformation by creating multi-year product roadmaps for integration and modernization projects")

# Tyson Foods
p = doc.add_paragraph()
run = p.add_run("Business Analyst/Software Engineer")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Gill Sans MT'
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0

p = doc.add_paragraph()
run = p.add_run("Tyson Foods")
run.font.color.rgb = DUSTY_BLUE
run.font.size = Pt(10)
run.font.name = 'Gill Sans MT'
run = p.add_run(" | 7/2010 - 6/2013")
run.font.size = Pt(10)
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing = 1.0

add_simple_bullet("Acted as key liaison between business stakeholders across sales, manufacturing, and accounting")
add_simple_bullet("Developed integrations that enhanced data flow between critical business systems")

# CF Industries
p = doc.add_paragraph()
run = p.add_run("Software Developer")
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Gill Sans MT'
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0

p = doc.add_paragraph()
run = p.add_run("CF Industries")
run.font.color.rgb = DUSTY_BLUE
run.font.size = Pt(10)
run.font.name = 'Gill Sans MT'
run = p.add_run(" | 4/2001 - 6/2010")
run.font.size = Pt(10)
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing = 1.0

add_simple_bullet("Engineered robust software solutions and integrations for enterprise manufacturing applications")

# ========== EDUCATION ==========
add_section_header("Education")

p = doc.add_paragraph()
run = p.add_run("Masters in Management")
run.bold = True
run.font.size = Pt(10.5)
run.font.name = 'Calibri'
run = p.add_run(" | Bellevue University, Bellevue, Nebraska")
run.font.size = Pt(10.5)
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(2)
p.paragraph_format.line_spacing = 1.0

p = doc.add_paragraph()
run = p.add_run("BS in Computer Science and Mathematics")
run.bold = True
run.font.size = Pt(10.5)
run.font.name = 'Calibri'
run = p.add_run(" | Morningside University, Sioux City, Iowa")
run.font.size = Pt(10.5)
run.font.name = 'Calibri'

# Save
output_path = r"g:\My Drive\GitHub\ExecutiveAssistant\Resume\output\Jorgensen Resume - Datamundi Insurance SME AI Evaluation 2026-01.docx"
doc.save(output_path)
print(f"Created: {output_path}")

# Convert to PDF
try:
    import win32com.client
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    pdf_path = output_path.replace('.docx', '.pdf')
    doc_obj = word.Documents.Open(output_path)
    doc_obj.SaveAs(pdf_path, FileFormat=17)
    doc_obj.Close()
    word.Quit()
    print(f"Created: {pdf_path}")
except Exception as e:
    print(f"PDF conversion failed: {e}")
    print("The .docx file was created successfully. Export to PDF manually from Word.")
